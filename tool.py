#!/usr/bin/env python3
# -*- coding: utf-8 -*-

'根据工程路径生成对应的软著用 doc 文档'

import sys

__author__ = 'Chuck Lin'

import os
import docx
from docx.shared import Pt

# 生成的 word 文档路径
_DOC_PATH = sys.argv[1]
# 项目工程路径
_PATHS = sys.argv[2].split(',')
# 页眉名称
_HEADER_NAME = sys.argv[3]

_G_DOC = docx.Document()
style = _G_DOC.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)
_G_DOC.sections[0].header.paragraphs[0].text = _HEADER_NAME
_INVALID_START_STR = ('*', '/**', '/**', ' * ', ' */')
g_doc_paragraph = _G_DOC.add_paragraph(style='Body Text')


def _rw_2_word(path):
    # 如果是 .java 后缀的文件则进行读取并写入 doc。
    with open(path, 'r', encoding='utf-8') as file:
        # line 中的内容凡事以非法字符串开头的都舍弃。
        valid_lines = [line for line in file.readlines() if _line_validator(line)]
    for valid_line in valid_lines:
        g_doc_paragraph.add_run(valid_line)
    _G_DOC.save(_DOC_PATH)


def _line_validator(line):
    for invalid in _INVALID_START_STR:
        if line.startswith(invalid):
            return False
    return True


def _listdir(path):
    for dirpath, dirname, files in os.walk(path):
        for file in files:
            file = os.path.join(dirpath, file)
            if os.path.splitext(file)[1] == '.java':
                _rw_2_word(file)


if __name__ == '__main__':
    for path in _PATHS:
        _listdir(path)
